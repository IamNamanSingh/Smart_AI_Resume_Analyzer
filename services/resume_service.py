import re
import io
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import PyPDF2
from core.database import get_database_connection

class ResumeAnalyzer:
    def __init__(self):
        self.document_types = {
            'resume': [
                'experience', 'education', 'skills', 'work', 'project', 'objective',
                'summary', 'employment', 'qualification', 'achievements'
            ],
            'marksheet': [
                'grade', 'marks', 'score', 'semester', 'cgpa', 'sgpa', 'examination',
                'result', 'academic year', 'percentage'
            ],
            'certificate': [
                'certificate', 'certification', 'awarded', 'completed', 'achievement',
                'training', 'course completion', 'qualified'
            ],
            'id_card': [
                'id card', 'identity', 'student id', 'employee id', 'valid until',
                'date of issue', 'identification'
            ]
        }
        
    def detect_document_type(self, text):
        text = text.lower()
        scores = {}
        
        for doc_type, keywords in self.document_types.items():
            matches = sum(1 for keyword in keywords if keyword in text)
            density = matches / len(keywords)
            frequency = matches / (len(text.split()) + 1)
            scores[doc_type] = (density * 0.7) + (frequency * 0.3)
        
        best_match = max(scores.items(), key=lambda x: x[1])
        return best_match[0] if best_match[1] > 0.15 else 'unknown'
        
    def calculate_keyword_match(self, resume_text, required_skills):
        resume_text = resume_text.lower()
        found_skills = []
        missing_skills = []
        
        for skill in required_skills:
            skill_lower = skill.lower()
            if skill_lower in resume_text:
                found_skills.append(skill)
            elif any(skill_lower in phrase for phrase in resume_text.split('.')):
                found_skills.append(skill)
            else:
                missing_skills.append(skill)
                
        match_score = (len(found_skills) / len(required_skills)) * 100 if required_skills else 0
        return {
            'score': match_score,
            'found_skills': found_skills,
            'missing_skills': missing_skills
        }
        
    def check_resume_sections(self, text):
        text = text.lower()
        essential_sections = {
            'contact': ['email', 'phone', 'address', 'linkedin'],
            'education': ['education', 'university', 'college', 'degree', 'academic'],
            'experience': ['experience', 'work', 'employment', 'job', 'internship'],
            'skills': ['skills', 'technologies', 'tools', 'proficiencies', 'expertise']
        }
        
        section_scores = {}
        for section, keywords in essential_sections.items():
            found = sum(1 for keyword in keywords if keyword in text)
            section_scores[section] = min(25, (found / len(keywords)) * 25)
            
        return sum(section_scores.values())
        
    def check_formatting(self, text):
        lines = text.split('\n')
        score = 100
        deductions = []
        
        if len(text) < 300:
            score -= 30
            deductions.append("Resume is too short")
            
        if not any(line.isupper() for line in lines):
            score -= 20
            deductions.append("No clear section headers found")
            
        if not any(line.strip().startswith(('•', '-', '*', '→')) for line in lines):
            score -= 20
            deductions.append("No bullet points found for listing details")
            
        if any(len(line.strip()) == 0 and len(next_line.strip()) == 0 
               for line, next_line in zip(lines[:-1], lines[1:])):
            score -= 15
            deductions.append("Inconsistent spacing between sections")
            
        contact_patterns = [
            r'\b[\w\.-]+@[\w\.-]+\.\w+\b',
            r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b',
            r'linkedin\.com/\w+',
        ]
        if not any(re.search(pattern, text) for pattern in contact_patterns):
            score -= 15
            deductions.append("Missing or improperly formatted contact information")
            
        return max(0, score), deductions
        
    def extract_personal_info(self, text):
        email_pattern = r'[\w\.-]+@[\w\.-]+\.\w+'
        phone_pattern = r'(\+\d{1,3}[-.]?)?\s*\(?\d{3}\)?[-.]?\s*\d{3}[-.]?\s*\d{4}'
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        github_pattern = r'github\.com/[\w-]+'
        
        email = re.search(email_pattern, text)
        phone = re.search(phone_pattern, text)
        linkedin = re.search(linkedin_pattern, text)
        github = re.search(github_pattern, text)
        
        name = text.split('\n')[0].strip()
        
        return {
            'name': name if len(name) > 0 else 'Unknown',
            'email': email.group(0) if email else '',
            'phone': phone.group(0) if phone else '',
            'linkedin': linkedin.group(0) if linkedin else '',
            'github': github.group(0) if github else '',
            'portfolio': ''
        }

    def extract_education(self, text):
        education = []
        lines = text.split('\n')
        education_keywords = [
            'education', 'academic', 'qualification', 'degree', 'university', 'college',
            'school', 'institute', 'certification', 'diploma', 'bachelor', 'master',
            'phd', 'b.tech', 'm.tech', 'b.e', 'm.e', 'b.sc', 'm.sc','bca', 'mca', 'b.com',
            'm.com', 'b.cs-it', 'imca', 'bba', 'mba', 'honors', 'scholarship'
        ]
        in_education_section = False
        current_entry = []

        for line in lines:
            line = line.strip()
            if any(keyword.lower() in line.lower() for keyword in education_keywords):
                if not any(keyword.lower() == line.lower() for keyword in education_keywords):
                    current_entry.append(line)
                in_education_section = True
                continue
            
            if in_education_section:
                if line and any(keyword.lower() in line.lower() for keyword in self.document_types['resume']):
                    if not any(edu_key.lower() in line.lower() for edu_key in education_keywords):
                        in_education_section = False
                        if current_entry:
                            education.append(' '.join(current_entry))
                            current_entry = []
                        continue
                
                if line:
                    current_entry.append(line)
                elif current_entry:
                    education.append(' '.join(current_entry))
                    current_entry = []
        
        if current_entry:
            education.append(' '.join(current_entry))
        return education

    def extract_experience(self, text):
        experience = []
        lines = text.split('\n')
        experience_keywords = [
            'experience', 'employment', 'work history', 'professional experience',
            'work experience', 'career history', 'professional background',
            'employment history', 'job history', 'positions held',
            'job title', 'job responsibilities', 'job description', 'job summary'
        ]
        in_experience_section = False
        current_entry = []

        for line in lines:
            line = line.strip()
            if any(keyword.lower() in line.lower() for keyword in experience_keywords):
                if not any(keyword.lower() == line.lower() for keyword in experience_keywords):
                    current_entry.append(line)
                in_experience_section = True
                continue
            
            if in_experience_section:
                if line and any(keyword.lower() in line.lower() for keyword in self.document_types['resume']):
                    if not any(exp_key.lower() in line.lower() for exp_key in experience_keywords):
                        in_experience_section = False
                        if current_entry:
                            experience.append(' '.join(current_entry))
                            current_entry = []
                        continue
                
                if line:
                    current_entry.append(line)
                elif current_entry:
                    experience.append(' '.join(current_entry))
                    current_entry = []
        
        if current_entry:
            experience.append(' '.join(current_entry))
        return experience

    def extract_projects(self, text):
        projects = []
        lines = text.split('\n')
        project_keywords = [
            'projects', 'personal projects', 'academic projects', 'key projects',
            'major projects', 'professional projects', 'project experience',
            'relevant projects', 'featured projects','latest projects',
            'top projects'
        ]
        in_project_section = False
        current_entry = []

        for line in lines:
            line = line.strip()
            if any(keyword.lower() in line.lower() for keyword in project_keywords):
                if not any(keyword.lower() == line.lower() for keyword in project_keywords):
                    current_entry.append(line)
                in_project_section = True
                continue
            
            if in_project_section:
                if line and any(keyword.lower() in line.lower() for keyword in self.document_types['resume']):
                    if not any(proj_key.lower() in line.lower() for proj_key in project_keywords):
                        in_project_section = False
                        if current_entry:
                            projects.append(' '.join(current_entry))
                            current_entry = []
                        continue
                
                if line:
                    current_entry.append(line)
                elif current_entry:
                    projects.append(' '.join(current_entry))
                    current_entry = []
        
        if current_entry:
            projects.append(' '.join(current_entry))
        return projects

    def extract_skills(self, text):
        skills = set()
        lines = text.split('\n')
        skills_keywords = [
            'skills', 'technical skills', 'competencies', 'expertise',
            'core competencies', 'professional skills', 'key skills',
            'technical expertise', 'proficiencies', 'qualifications',
            'top skills', 'key skill', 'major skill', 'personal skill',
            'soft skills', 'soft skill'
        ]
        in_skills_section = False
        current_entry = []
        separators = [',', '•', '|', '/', '\\', '·', '>', '-', '–', '―']

        for line in lines:
            line = line.strip()
            if any(keyword.lower() in line.lower() for keyword in skills_keywords):
                if not any(keyword.lower() == line.lower() for keyword in skills_keywords):
                    current_entry.append(line)
                in_skills_section = True
                continue
            
            if in_skills_section:
                if line and any(keyword.lower() in line.lower() for keyword in self.document_types['resume']):
                    if not any(skill_key.lower() in line.lower() for skill_key in skills_keywords):
                        in_skills_section = False
                        if current_entry:
                            text_to_process = ' '.join(current_entry)
                            for separator in separators:
                                if separator in text_to_process:
                                    skills.update(skill.strip() for skill in text_to_process.split(separator) if skill.strip())
                            current_entry = []
                        continue
                
                if line:
                    current_entry.append(line)
                elif current_entry:
                    text_to_process = ' '.join(current_entry)
                    for separator in separators:
                        if separator in text_to_process:
                            skills.update(skill.strip() for skill in text_to_process.split(separator) if skill.strip())
                    current_entry = []
        
        if current_entry:
            text_to_process = ' '.join(current_entry)
            for separator in separators:
                if separator in text_to_process:
                    skills.update(skill.strip() for skill in text_to_process.split(separator) if skill.strip())
        return list(skills)

    def extract_summary(self, text):
        summary = []
        lines = text.split('\n')
        summary_keywords = [
            'summary', 'professional summary', 'career summary', 'objective',
            'career objective', 'professional objective', 'about me', 'profile',
            'professional profile', 'career profile', 'overview'
        ]
        in_summary_section = False
        current_entry = []

        start_index = 0
        while start_index < min(10, len(lines)) and not lines[start_index].strip():
            start_index += 1

        first_lines = []
        lines_checked = 0
        for line in lines[start_index:]:
            if line.strip():
                first_lines.append(line.strip())
                lines_checked += 1
                if lines_checked >= 5:
                    break

        if first_lines and not any(keyword in first_lines[0].lower() for keyword in summary_keywords):
            potential_summary = ' '.join(first_lines)
            if len(potential_summary.split()) > 10:
                if not re.search(r'\b(?:email|phone|address|tel|mobile|linkedin)\b', potential_summary.lower()):
                    summary.append(potential_summary)

        for line in lines:
            line = line.strip()
            if any(keyword.lower() in line.lower() for keyword in summary_keywords):
                if not any(keyword.lower() == line.lower() for keyword in summary_keywords):
                    current_entry.append(line)
                in_summary_section = True
                continue
            
            if in_summary_section:
                if line and any(keyword.lower() in line.lower() for keyword in self.document_types['resume']):
                    if not any(sum_key.lower() in line.lower() for sum_key in summary_keywords):
                        in_summary_section = False
                        if current_entry:
                            summary.append(' '.join(current_entry))
                            current_entry = []
                        continue
                
                if line:
                    current_entry.append(line)
                elif current_entry:
                    summary.append(' '.join(current_entry))
                    current_entry = []
        
        if current_entry:
            summary.append(' '.join(current_entry))
        return ' '.join(summary) if summary else ''

    def analyze_resume(self, resume_data, job_requirements):
        """Standard analysis of a resume text dataset"""
        try:
            text = resume_data.get('raw_text', '')
            personal_info = self.extract_personal_info(text)
            doc_type = self.detect_document_type(text)
            
            if doc_type != 'resume':
                return {
                    'ats_score': 0,
                    'document_type': doc_type,
                    'keyword_match': {'score': 0, 'found_skills': [], 'missing_skills': []},
                    'section_score': 0,
                    'format_score': 0,
                    'suggestions': [f"This document appears to be a {doc_type}. Please upload a standard resume."]
                }
                
            required_skills = job_requirements.get('required_skills', [])
            keyword_match = self.calculate_keyword_match(text, required_skills)
            
            education = self.extract_education(text)
            experience = self.extract_experience(text)
            projects = self.extract_projects(text)
            skills = list(self.extract_skills(text))
            summary = self.extract_summary(text)
            
            section_score = self.check_resume_sections(text)
            format_score, format_deductions = self.check_formatting(text)
            
            contact_suggestions = []
            if not personal_info.get('email'): contact_suggestions.append("Add your email address")
            if not personal_info.get('phone'): contact_suggestions.append("Add your phone number")
            if not personal_info.get('linkedin'): contact_suggestions.append("Add your LinkedIn URL")
            
            summary_suggestions = []
            if not summary:
                summary_suggestions.append("Add a professional summary section")
            elif len(summary.split()) < 30:
                summary_suggestions.append("Expand your summary to better highlight your value")
            
            skills_suggestions = []
            if not skills:
                skills_suggestions.append("Add a dedicated skills section")
            if len(skills) < 5:
                skills_suggestions.append("List more relevant skills")
            
            experience_suggestions = []
            if not experience:
                experience_suggestions.append("Add your work experience section")
            
            education_suggestions = []
            if not education:
                education_suggestions.append("Add your educational background")
            
            format_suggestions = list(format_deductions)
            
            contact_score = 100 - (len(contact_suggestions) * 25)
            summary_score = 100 - (len(summary_suggestions) * 33)
            skills_score = keyword_match['score']
            experience_score = 100 - (len(experience_suggestions) * 25)
            education_score = 100 - (len(education_suggestions) * 25)
            
            ats_score = (
                int(round(contact_score * 0.1)) +
                int(round(summary_score * 0.1)) +
                int(round(skills_score * 0.3)) +
                int(round(experience_score * 0.2)) +
                int(round(education_score * 0.1)) +
                int(round(format_score * 0.2))
            )
            
            suggestions = contact_suggestions + summary_suggestions + skills_suggestions + experience_suggestions + education_suggestions + format_suggestions
            if not suggestions:
                suggestions.append("Your resume is well-optimized for ATS systems")
                
            return {
                **personal_info,
                'ats_score': ats_score,
                'document_type': 'resume',
                'keyword_match': keyword_match,
                'section_score': section_score,
                'format_score': format_score,
                'education': education,
                'experience': experience,
                'projects': projects,
                'skills': skills,
                'summary': summary,
                'suggestions': suggestions,
                'contact_suggestions': contact_suggestions,
                'summary_suggestions': summary_suggestions,
                'skills_suggestions': skills_suggestions,
                'experience_suggestions': experience_suggestions,
                'education_suggestions': education_suggestions,
                'format_suggestions': format_suggestions,
                'section_scores': {
                    'contact': contact_score,
                    'summary': summary_score,
                    'skills': skills_score,
                    'experience': experience_score,
                    'education': education_score,
                    'format': format_score
                }
            }
        except Exception as e:
            return {
                'error': str(e),
                'ats_score': 0,
                'document_type': 'unknown',
                'suggestions': [f"Error during standard analysis: {str(e)}"]
            }


class ResumeBuilder:
    def __init__(self):
        pass

    def _format_list_items(self, items):
        if isinstance(items, str):
            return [item.strip() for item in items.split('\n') if item.strip()]
        elif isinstance(items, list):
            return [item.strip() for item in items if item and item.strip()]
        return []

    def generate_resume(self, data):
        """Generate a Word document file from resume builder form data"""
        try:
            doc = Document()
            template_name = data.get('template', 'Modern').lower()
            
            if template_name == 'professional':
                doc = self.build_professional_template(doc, data)
            elif template_name == 'minimal':
                doc = self.build_minimal_template(doc, data)
            elif template_name == 'creative':
                doc = self.build_creative_template(doc, data)
            else:
                doc = self.build_modern_template(doc, data)
                
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer
        except Exception as e:
            print(f"Error generating resume: {e}")
            raise

    def build_modern_template(self, doc, data):
        # Set up modern margins
        for section in doc.sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # Base font configurations
        styles = doc.styles
        normal_style = styles.add_style('Mod Normal', WD_STYLE_TYPE.PARAGRAPH) if 'Mod Normal' not in styles else styles['Mod Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(10)
        normal_style.font.color.rgb = RGBColor(44, 62, 80)
        
        # Add name header
        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_name.paragraph_format.space_before = Pt(10)
        p_name.paragraph_format.space_after = Pt(2)
        run_name = p_name.add_run(data['personal_info'].get('full_name', '').upper())
        run_name.font.name = 'Arial'
        run_name.font.size = Pt(22)
        run_name.bold = True
        run_name.font.color.rgb = RGBColor(41, 128, 185)

        # Contact info
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_contact.paragraph_format.space_after = Pt(10)
        p_contact.style = normal_style
        
        parts = []
        pi = data['personal_info']
        if pi.get('email'): parts.append(pi['email'])
        if pi.get('phone'): parts.append(pi['phone'])
        if pi.get('location'): parts.append(pi['location'])
        if pi.get('linkedin'): parts.append(pi['linkedin'])
        p_contact.add_run(" | ".join(parts))

        # Professional Summary
        if data.get('summary'):
            p_head = doc.add_paragraph()
            p_head.paragraph_format.space_before = Pt(12)
            run = p_head.add_run("PROFESSIONAL SUMMARY")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(41, 128, 185)
            
            p_sum = doc.add_paragraph(data['summary'], style=normal_style)
            p_sum.paragraph_format.left_indent = Inches(0.15)

        # Experience
        if data.get('experience'):
            p_head = doc.add_paragraph()
            p_head.paragraph_format.space_before = Pt(12)
            run = p_head.add_run("WORK EXPERIENCE")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(41, 128, 185)
            
            for exp in data['experience']:
                p_exp = doc.add_paragraph(style=normal_style)
                p_exp.paragraph_format.left_indent = Inches(0.15)
                p_exp.add_run(f"{exp.get('position', '')} at {exp.get('company', '')}").bold = True
                p_exp.add_run(f"  ({exp.get('start_date', '')} - {exp.get('end_date', '')})").italic = True
                
                if exp.get('description'):
                    doc.add_paragraph(exp['description'], style=normal_style).paragraph_format.left_indent = Inches(0.3)
                if exp.get('responsibilities'):
                    for resp in self._format_list_items(exp['responsibilities']):
                        bullet = doc.add_paragraph(style=normal_style)
                        bullet.paragraph_format.left_indent = Inches(0.4)
                        bullet.add_run("• " + resp)

        # Education
        if data.get('education'):
            p_head = doc.add_paragraph()
            p_head.paragraph_format.space_before = Pt(12)
            run = p_head.add_run("EDUCATION")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(41, 128, 185)
            
            for edu in data['education']:
                p_edu = doc.add_paragraph(style=normal_style)
                p_edu.paragraph_format.left_indent = Inches(0.15)
                p_edu.add_run(f"{edu.get('degree', '')} in {edu.get('field', '')}").bold = True
                p_edu.add_run(f" | {edu.get('school', '')} ({edu.get('graduation_date', '')})")
                if edu.get('gpa'):
                    p_edu.add_run(f"  GPA: {edu['gpa']}")

        # Skills
        if data.get('skills'):
            p_head = doc.add_paragraph()
            p_head.paragraph_format.space_before = Pt(12)
            run = p_head.add_run("SKILLS")
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(41, 128, 185)
            
            p_skills = doc.add_paragraph(style=normal_style)
            p_skills.paragraph_format.left_indent = Inches(0.15)
            
            skills_dict = data['skills']
            skills_parts = []
            if isinstance(skills_dict, dict):
                for k, v in skills_dict.items():
                    val_list = self._format_list_items(v)
                    if val_list:
                        skills_parts.append(f"{k.capitalize()}: {', '.join(val_list)}")
            elif isinstance(skills_dict, list):
                skills_parts.append(", ".join(skills_dict))
                
            p_skills.add_run(" | ".join(skills_parts))

        return doc

    def build_professional_template(self, doc, data):
        # Implementation similar to modern but uses different font/colors (e.g. Calibri, darkblue)
        return self.build_modern_template(doc, data)

    def build_minimal_template(self, doc, data):
        return self.build_modern_template(doc, data)

    def build_creative_template(self, doc, data):
        return self.build_modern_template(doc, data)


# Utility file text extractions
def extract_text_from_pdf(pdf_file):
    """Fallback standard PDF text extraction"""
    text = ""
    try:
        if hasattr(pdf_file, 'read'):
            file_content = pdf_file.read()
            pdf_file.seek(0)
        else:
            file_content = pdf_file
            
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Failed standard PDF text extraction: {e}")

def extract_text_from_docx(docx_file):
    """Standard DOCX text extraction"""
    try:
        doc = Document(docx_file)
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        return '\n'.join(full_text)
    except Exception as e:
        raise Exception(f"Failed DOCX text extraction: {e}")

def export_resumes_to_excel():
    """Generates excel dataset bytes representing all resumes in the database"""
    conn = get_database_connection()
    query = """
        SELECT
            rd.name, rd.email, rd.phone, rd.linkedin, rd.github, rd.portfolio,
            rd.summary, rd.target_role, rd.target_category,
            rd.education, rd.experience, rd.projects, rd.skills,
            ra.ats_score, ra.keyword_match_score, ra.format_score, ra.section_score,
            ra.missing_skills, ra.recommendations,
            rd.created_at
        FROM resume_data rd
        LEFT JOIN resume_analysis ra ON rd.id = ra.resume_id
    """
    try:
        df = pd.read_sql_query(query, conn)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Resume Data')
        return output.getvalue()
    except Exception as e:
        print(f"Error exporting to Excel: {e}")
        return None
    finally:
        conn.close()
